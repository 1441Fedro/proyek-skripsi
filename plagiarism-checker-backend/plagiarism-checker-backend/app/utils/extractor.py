import os
import zipfile
import patoolib
import pdfplumber
from pathlib import Path
from io import BytesIO
from docx import Document

# def extract_logic_text(file_path: str) -> str:
#     ext = Path(file_path).suffix.lower()
#     try:
#         if ext == ".pdf":
#             return extract_logic_from_pdf(file_path)
#         elif ext == ".docx":
#             return extract_logic_from_docx(file_path)
#         elif ext in [".zip", ".rar"]:
#             return extract_logic_from_archive(file_path)
#     except:
#         return None
#     return None

# from pathlib import Path

def extract_logic_text(file_path: str) -> str:
    ext = Path(file_path).suffix.lower()
    try:
        if ext == ".pdf":
            print(f"[INFO] Ekstraksi PDF: {file_path}")
            result = extract_logic_from_pdf(file_path)
        elif ext == ".docx":
            print(f"[INFO] Ekstraksi DOCX: {file_path}")
            result = extract_logic_from_docx(file_path)
        elif ext in [".zip", ".rar"]:
            print(f"[INFO] Ekstraksi arsip: {file_path}")
            result = extract_logic_from_archive(file_path)
        else:
            print(f"[WARN] Ekstensi tidak dikenali: {file_path}")
            result = ""
        return result
    except Exception as e:
        print(f"[ERROR] Gagal ekstrak file: {file_path}, error: {e}")
        return ""


def extract_full_text(file_path: str) -> str:
    ext = Path(file_path).suffix.lower()
    try:
        if ext == ".pdf":
            return extract_text_from_pdf(file_path)
        elif ext == ".docx":
            return extract_text_from_docx(file_path)
        elif ext in [".zip", ".rar"]:
            return extract_text_from_zip_or_rar(file_path)
    except:
        return None
    return None

# ========== PDF ==========
# def extract_logic_from_pdf(path):
#     with pdfplumber.open(path) as pdf:
#         text = ""
#         logic_section = False
#         for i, page in enumerate(pdf.pages):
#             if i < 2:  # Skip halaman 1 dan 2 (cover & listing)
#                 continue
#             page_text = page.extract_text() or ""
#             if "logika" in page_text.lower():
#                 logic_section = True
#             if logic_section:
#                 text += page_text[index:] + "\n"
#         print(text)
#         result = text.replace('\n', ' ').lower().strip()
#         return result if result else ""
    # return text.replace('\n', ' ').lower().strip()

def extract_logic_from_pdf(path):
    with pdfplumber.open(path) as pdf:
        text = ""
        logic_section = False

        for i, page in enumerate(pdf.pages):
            if i < 2:
                continue  # Lewati cover & daftar isi

            page_text = page.extract_text() or ""

            if "logika" in page_text.lower():
                logic_section = True

            if logic_section:
                text += page_text + "\n"

        # Pastikan selalu return string
        result = text.replace('\n', ' ').lower().strip()
        return result if result else ""

def extract_text_from_pdf(pdf_path: str) -> str:
    with pdfplumber.open(pdf_path) as pdf:
        return "\n".join(page.extract_text() or "" for page in pdf.pages)

# ========== DOCX ==========
def extract_logic_from_docx(path):
    doc = Document(path)
    text = ""
    logic_section = False
    for para in doc.paragraphs:
        if "logika" in para.text.lower():
            logic_section = True
        if logic_section:
            text += para.text + "\n"
    return text.replace('\n', ' ').lower().strip()
    # text.strip()

def extract_text_from_docx(docx_path: str) -> str:
    doc = Document(docx_path)
    return "\n".join(para.text for para in doc.paragraphs)

# ========== ARCHIVES ==========
def extract_logic_from_archive(file_path: str) -> str:
    extracted_text = ""
    temp_folder = "temp_logic_extract"

    if file_path.endswith(".zip"):
        with zipfile.ZipFile(file_path, "r") as zip_ref:
            zip_ref.extractall(temp_folder)
    elif file_path.endswith(".rar"):
        patoolib.extract_archive(file_path, outdir=temp_folder)

    for root, dirs, files in os.walk(temp_folder):
        for file in files:
            path = os.path.join(root, file)
            ext = Path(file).suffix.lower()
            if ext in [".pdf", ".docx"]:
                text = extract_logic_text(path)
                if text:
                    extracted_text += text + "\n"
    # Optional: bersihkan temp folder
    # shutil.rmtree(temp_folder)
    return extracted_text.replace('\n', ' ').lower().strip()

def extract_text_from_zip_or_rar(file_path: str) -> str:
    extracted_text = ""
    temp_folder = "temp_full_extract"

    if file_path.endswith('.zip'):
        with zipfile.ZipFile(file_path, 'r') as zip_ref:
            zip_ref.extractall(temp_folder)
    elif file_path.endswith('.rar'):
        patoolib.extract_archive(file_path, outdir=temp_folder)

    for root, _, files in os.walk(temp_folder):
        for file in files:
            path = os.path.join(root, file)
            ext = Path(file).suffix.lower()
            if ext == ".pdf":
                extracted_text += extract_text_from_pdf(path) + "\n"
            elif ext == ".docx":
                extracted_text += extract_text_from_docx(path) + "\n"

    return extracted_text.strip()